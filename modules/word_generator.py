from docx import Document
from docx.shared import Pt, RGBColor
import json
import os
from datetime import datetime

class WordGenerator:
    def __init__(self, templates_dir="templates"):
        self.templates_dir = templates_dir
        self.config = self.load_config()
    
    def load_config(self):
        """Carrega configuração de templates"""
        config_path = os.path.join(self.templates_dir, "config.json")
        with open(config_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def get_template_info(self, template_id):
        """Retorna informações do template"""
        for template in self.config['templates']:
            if template['id'] == template_id:
                return template
        return None
    
    def generate_document(self, template_id, candidate_data):
        """
        Gera documento Word preenchido com dados do candidato
        
        Args:
            template_id: ID do template (ex: 'cv', 'carta')
            candidate_data: Dicionário com dados do candidato
        
        Returns:
            Caminho do arquivo gerado
        """
        # Obter informações do template
        template_info = self.get_template_info(template_id)
        if not template_info:
            raise ValueError(f"Template '{template_id}' não encontrado")
        
        # Carregar template
        template_path = os.path.join(
            self.templates_dir,
            template_info['arquivo']
        )
        doc = Document(template_path)
        
        # Substituir placeholders
        placeholders = template_info['placeholders']
        for placeholder, field_name in placeholders.items():
            value = candidate_data.get(field_name, "")
            self._replace_text_in_document(doc, placeholder, str(value))
        
        # Salvar documento
        output_filename = f"{candidate_data['nome']}_{template_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join("/tmp", output_filename)
        doc.save(output_path)
        
        return output_path, output_filename
    
    def _replace_text_in_document(self, doc, placeholder, replacement):
        """Substitui placeholder no documento"""
        # Substituir em parágrafos
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                self._replace_text_in_paragraph(paragraph, placeholder, replacement)
        
        # Substituir em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            self._replace_text_in_paragraph(paragraph, placeholder, replacement)
    
    def _replace_text_in_paragraph(self, paragraph, placeholder, replacement):
        """Substitui placeholder em um parágrafo específico"""
        if placeholder in paragraph.text:
            # Limpar parágrafo
            for run in paragraph.runs:
                run.text = run.text.replace(placeholder, replacement)
    
    def list_templates(self):
        """Retorna lista de templates disponíveis"""
        return [
            {
                "id": t['id'],
                "nome": t['nome']
            }
            for t in self.config['templates']
        ]